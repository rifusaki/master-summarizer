"""
DOCX document parser.

Extracts text, tables, and images from Word documents, preserving
structural metadata (headings, sections, paragraph indices) for
downstream chunking and provenance tracking.
"""

from __future__ import annotations

import base64
import io
import logging
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image

from src.models import (
    ArtifactType,
    DocumentParseResult,
    NormalizedArtifact,
    SourceLocation,
)

logger = logging.getLogger(__name__)

# Heading style name prefixes
HEADING_STYLES = {"Heading", "heading", "Title", "Título", "Encabezado"}


def _is_heading(paragraph: Paragraph) -> bool:
    """Check if a paragraph is a heading."""
    style_name = paragraph.style.name if paragraph.style else ""
    return any(style_name.startswith(h) for h in HEADING_STYLES)


def _heading_level(paragraph: Paragraph) -> int:
    """Extract heading level (1-9) from a heading paragraph."""
    style_name = paragraph.style.name if paragraph.style else ""
    # Try to extract number from style name like "Heading 2"
    for part in style_name.split():
        if part.isdigit():
            return int(part)
    # Title = level 1
    if "title" in style_name.lower() or "título" in style_name.lower():
        return 1
    return 1


def _table_to_dict(table: Table) -> list[dict[str, Any]]:
    """Convert a DOCX table to a list of row dicts."""
    rows_data: list[dict[str, Any]] = []
    headers: list[str] = []

    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]

        if row_idx == 0:
            # Use first row as headers
            headers = cells if any(cells) else [f"col_{i}" for i in range(len(cells))]
            continue

        if headers:
            row_dict = {h: v for h, v in zip(headers, cells)}
        else:
            row_dict = {f"col_{i}": v for i, v in enumerate(cells)}
        rows_data.append(row_dict)

    return rows_data


def _table_to_text(table: Table) -> str:
    """Convert a DOCX table to a readable text representation."""
    lines: list[str] = []
    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        lines.append(" | ".join(cells))
        if row_idx == 0:
            lines.append("-" * len(lines[-1]))
    return "\n".join(lines)


def _extract_images(
    doc: DocxDocument, source_file: str, output_dir: Path
) -> list[NormalizedArtifact]:
    """Extract all embedded images from a DOCX document."""
    artifacts: list[NormalizedArtifact] = []
    output_dir.mkdir(parents=True, exist_ok=True)
    image_idx = 0

    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            try:
                image_data = rel.target_part.blob
                # Determine format
                content_type = rel.target_part.content_type
                ext = "png"
                if "jpeg" in content_type or "jpg" in content_type:
                    ext = "jpg"
                elif "gif" in content_type:
                    ext = "gif"
                elif "tiff" in content_type:
                    ext = "tiff"

                # Save image file
                img_filename = f"{Path(source_file).stem}_img_{image_idx:04d}.{ext}"
                img_path = output_dir / img_filename
                img_path.write_bytes(image_data)

                # Get image dimensions
                img = Image.open(io.BytesIO(image_data))
                width, height = img.size

                # Base64 encode for multimodal processing
                b64 = base64.b64encode(image_data).decode("utf-8")

                artifact = NormalizedArtifact(
                    artifact_type=ArtifactType.IMAGE,
                    content=f"[Image: {img_filename}, {width}x{height}px, {ext}]",
                    source=SourceLocation(
                        source_file=source_file,
                        image_index=image_idx,
                    ),
                    image_path=str(img_path),
                    image_base64=b64,
                    metadata={
                        "width": width,
                        "height": height,
                        "format": ext,
                        "size_bytes": len(image_data),
                        "rel_id": rel_id,
                    },
                )
                artifacts.append(artifact)
                image_idx += 1

            except Exception as exc:
                logger.warning(
                    "Failed to extract image %s from %s: %s",
                    rel_id,
                    source_file,
                    exc,
                )

    logger.info("Extracted %d images from %s", len(artifacts), source_file)
    return artifacts


def parse_docx(
    filepath: Path, images_output_dir: Path | None = None
) -> DocumentParseResult:
    """
    Parse a DOCX file into normalized artifacts.

    Extracts text blocks (preserving heading hierarchy), tables,
    and embedded images. No summarization is performed.

    Args:
        filepath: Path to the DOCX file.
        images_output_dir: Directory to save extracted images.

    Returns:
        DocumentParseResult with all extracted artifacts.
    """
    source_file = filepath.name
    logger.info("Parsing DOCX: %s", source_file)

    doc = DocxDocument(str(filepath))
    artifacts: list[NormalizedArtifact] = []
    heading_structure: list[dict[str, Any]] = []
    current_heading_path: list[str] = []
    current_section = ""
    paragraph_idx = 0
    table_idx = 0
    total_text_length = 0

    # Process document body elements in order
    # We need to iterate through the document body to maintain ordering
    # between paragraphs and tables
    body = doc.element.body
    para_iter = iter(doc.paragraphs)
    table_iter = iter(doc.tables)

    current_para = next(para_iter, None)
    current_table = next(table_iter, None)

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p" and current_para is not None:
            paragraph = current_para
            current_para = next(para_iter, None)
            text = paragraph.text.strip()

            if not text:
                continue

            if _is_heading(paragraph):
                level = _heading_level(paragraph)
                # Update heading path
                while len(current_heading_path) >= level:
                    current_heading_path.pop()
                current_heading_path.append(text)
                current_section = text

                heading_structure.append(
                    {
                        "level": level,
                        "text": text,
                        "path": list(current_heading_path),
                        "paragraph_index": paragraph_idx,
                    }
                )

                artifact = NormalizedArtifact(
                    artifact_type=ArtifactType.TEXT,
                    content=text,
                    raw_content=text,
                    source=SourceLocation(
                        source_file=source_file,
                        section=current_section,
                        heading_path=list(current_heading_path),
                        paragraph_index=paragraph_idx,
                    ),
                    metadata={"is_heading": True, "heading_level": level},
                )
            else:
                artifact = NormalizedArtifact(
                    artifact_type=ArtifactType.TEXT,
                    content=text,
                    raw_content=text,
                    source=SourceLocation(
                        source_file=source_file,
                        section=current_section,
                        heading_path=list(current_heading_path),
                        paragraph_index=paragraph_idx,
                    ),
                    metadata={"is_heading": False},
                )
                total_text_length += len(text)

            artifacts.append(artifact)
            paragraph_idx += 1

        elif tag == "tbl" and current_table is not None:
            table = current_table
            current_table = next(table_iter, None)

            table_data = _table_to_dict(table)
            table_text = _table_to_text(table)

            artifact = NormalizedArtifact(
                artifact_type=ArtifactType.TABLE,
                content=table_text,
                raw_content=table_text,
                table_data=table_data,
                source=SourceLocation(
                    source_file=source_file,
                    section=current_section,
                    heading_path=list(current_heading_path),
                    table_index=table_idx,
                ),
                metadata={
                    "rows": len(table_data),
                    "columns": len(table_data[0]) if table_data else 0,
                },
            )
            artifacts.append(artifact)
            total_text_length += len(table_text)
            table_idx += 1

    # Extract images
    if images_output_dir is None:
        from src.config import PREPROCESSED_DIR

        images_output_dir = PREPROCESSED_DIR / "images"

    image_artifacts = _extract_images(doc, source_file, images_output_dir)
    artifacts.extend(image_artifacts)

    result = DocumentParseResult(
        source_file=source_file,
        title=current_heading_path[0] if current_heading_path else source_file,
        artifacts=artifacts,
        heading_structure=heading_structure,
        total_text_length=total_text_length,
        total_tables=table_idx,
        total_images=len(image_artifacts),
    )

    logger.info(
        "Parsed %s: %d artifacts (%d text, %d tables, %d images), %d chars",
        source_file,
        len(artifacts),
        paragraph_idx,
        table_idx,
        len(image_artifacts),
        total_text_length,
    )

    return result
